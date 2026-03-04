#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
auto_q.py — unified learner for mixed enterprise questionnaires

Implements Category 1–5 learnings:

  • Cat 1: PDF questionnaires with header/th cells, instruction rows, and checkbox clusters
  • Cat 2: Legal/contract PDFs (BAA, Amendments, COI) — skip learn
  • Cat 3: True questionnaire DOCX with Q/A tables — parse cautiously, reject placeholders/checkbox-only
  • Cat 4: Enterprise Technology Surveys (customer standards + vendor responses) —
           skip standards; learn only vendor responses; normalize technical facts
  • Cat 5: Checkbox-heavy forms with placeholders —
           infer booleans only when unambiguous; prefer narrative; drop placeholders

Outputs
  • kb.learned.json   — accepted Q/A items (with normalization metadata)
  • learn_report.json — skip/accept accounting with examples
  • tech_profile.json — consolidated technical profile (ports, HA, encryption, latency, RTO, auth, support access)
"""

import argparse
import json
import os
import re
import sys
import warnings
from collections import defaultdict
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple, Any, Iterable

try:
    from docx import Document
except Exception:
    Document = None

try:
    import pandas as pd
except Exception:
    pd = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

# Silence openpyxl noise
warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')

PLACEHOLDER_PATTERNS = [
    r"^\s*click here to enter text\.?\s*$",
    r"^\s*n/?a\s*$",
    r"^\s*not applicable\s*$",
    r"^\s*discussion\s*$",
    r"^\s*see attached\s*$",
    r"^\s*choose an item\.?\s*$",
    r"^\s*—\s*$",
    r"^_+$",
]
PLACEHOLDER_RE = re.compile("|".join(PLACEHOLDER_PATTERNS), re.IGNORECASE)
GUID_RE = re.compile(r"[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}", re.I)

LEGAL_TRIGGERS = [
    "business associate agreement","baa","first amendment","amendment","appendix e",
    "certificate of liability insurance","acord","terms and conditions","limitation of liability",
]
STANDARD_TRIGGERS = [
    "standard is to","baycare standard","standards:","general standards","baycare uses",
    "please review","please respond","if this section doesn’t apply","vendor sign-off","sign-off",
    "all traffic","san-connected","all production servers will be patched","anti-virus",
    "device standards:","voice standards","unified communications (uc) standards",
]
YESNO_TOKENS = {"yes","no"}

def normtxt(x: Any) -> str:
    if x is None: return ""
    return re.sub(r"\s+"," ",str(x)).strip()

def looks_like_placeholder(s: str) -> bool:
    s = normtxt(s)
    if not s: return True
    if PLACEHOLDER_RE.search(s): return True
    if GUID_RE.search(s): return True
    if s.lower() in YESNO_TOKENS: return True
    if len(re.findall(r"yes|no|☐|\[\s*\]", s, flags=re.I)) >= 2 and len(s) <= 12: return True
    return False

def is_customer_standard_block(s: str) -> bool:
    s = s.lower()
    return any(t in s for t in STANDARD_TRIGGERS)

def is_legal_contract_text(s: str) -> bool:
    s = s.lower()
    return any(t in s for t in LEGAL_TRIGGERS)

def is_blank_survey_pdf(text: str) -> bool:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    num_prompts = sum(1 for l in lines if re.match(r"^\d{1,3}[\).] ", l))
    yn = len(re.findall(r"\bYes\b|\bNo\b|☐", text, flags=re.I))
    long_lines = sum(1 for l in lines if len(l) > 120)
    return (num_prompts >= 8 and yn >= 10 and long_lines < 5)

@dataclass
class LearnItem:
    source: str
    question: str
    answer: str
    section: Optional[str] = None
    normalized: Dict[str, Any] = field(default_factory=dict)

@dataclass
class LearnReport:
    accepted: int = 0
    skipped: Dict[str,int] = field(default_factory=lambda: defaultdict(int))
    examples: Dict[str, List[Dict[str,str]]] = field(default_factory=lambda: defaultdict(list))
    def skip(self, reason: str, file: str, extra: Optional[Dict[str,str]] = None):
        self.skipped[reason] += 1
        if len(self.examples[reason]) < 8:
            ex = {"file": os.path.basename(file)}
            if extra: ex.update(extra)
            self.examples[reason].append(ex)
    def accept(self): self.accepted += 1

class TechProfile:
    def __init__(self):
        self.ports_tcp: List[str] = []
        self.ports_udp: List[str] = []
        self.rto_hours_max: Optional[float] = None
        self.max_client_latency_ms: Optional[int] = None
        self.auth: Optional[str] = None
        self.support_access: Optional[str] = None
        self.ha_db: Optional[str] = None
        self.ha_app: Optional[str] = None
        self.encryption_at_rest: Optional[str] = None
        self.encryption_in_transit: Optional[str] = None
    def as_dict(self) -> Dict[str,Any]:
        return {
            "ports":{"tcp": self.ports_tcp, "udp": self.ports_udp},
            "rto_hours_max": self.rto_hours_max,
            "max_client_latency_ms": self.max_client_latency_ms,
            "auth": self.auth,
            "support_access": self.support_access,
            "ha": {"db": self.ha_db, "app": self.ha_app},
            "encryption": {"at_rest": self.encryption_at_rest, "in_transit": self.encryption_in_transit},
        }
    def merge_ports(self, answer: str):
        s = answer.lower()
        udp_matches = re.findall(r"udp\s*:?\s*([0-9,\-\s]+)", s)
        for m in udp_matches:
            for part in re.split(r"[,;\s]+", m.strip()):
                if not part: continue
                if part.isdigit():
                    val = int(part)
                    if 1 <= val <= 65535 and str(val) not in self.ports_udp:
                        self.ports_udp.append(str(val))
                elif re.match(r"^\d{1,5}-\d{1,5}$", part):
                    lo, hi = map(int, part.split("-"))
                    if 1 <= lo <= 65535 and 1 <= hi <= 65535 and lo <= hi and part not in self.ports_udp:
                        self.ports_udp.append(part)
        cleaned = re.sub(r"udp\s*:?[^,;\n]+", "", s, flags=re.I)
        for token in re.split(r"[,;\s]+", cleaned):
            token = token.strip()
            if not token: continue
            if re.match(r"^\d+$", token):
                val = int(token)
                if 1 <= val <= 65535 and token not in self.ports_tcp:
                    self.ports_tcp.append(token)
            elif re.match(r"^\d{1,5}-\d{1,5}$", token):
                lo, hi = map(int, token.split("-"))
                if 1 <= lo <= 65535 and 1 <= hi <= 65535 and lo <= hi and token not in self.ports_tcp:
                    self.ports_tcp.append(token)
    def set_latency(self, text: str):
        m = re.search(r"(\d{1,3})\s*ms", text.lower())
        if m:
            val = int(m.group(1))
            self.max_client_latency_ms = val if (self.max_client_latency_ms is None or val < self.max_client_latency_ms) else self.max_client_latency_ms
    def set_rto(self, text: str):
        m = re.search(r"(\d+(?:\.\d+)?)\s*(?:to|–|-|—)\s*(\d+(?:\.\d+)?)\s*hour", text.lower())
        if m:
            hi = float(m.group(2))
            self.rto_hours_max = hi if (self.rto_hours_max is None or hi < self.rto_hours_max) else self.rto_hours_max
        else:
            m2 = re.search(r"(\d+(?:\.\d+)?)\s*hour", text.lower())
            if m2:
                val = float(m2.group(1))
                self.rto_hours_max = val if (self.rto_hours_max is None or val < self.rto_hours_max) else self.rto_hours_max
    def set_auth(self, text: str):
        if re.search(r"active directory|\bad\b|ldap|saml|single sign[- ]on|imprivata", text, flags=re.I):
            if re.search(r"active directory|\bAD\b", text, flags=re.I): self.auth = "Active Directory"
            elif re.search(r"saml|ldap", text, flags=re.I): self.auth = "Directory/SSO compatible"
    def set_support_access(self, text: str):
        if re.search(r"site[\-\s]?to[\-\s]?site\s+vpn|site\s*\-\s*site\s+vpn|s2s\s+vpn", text, flags=re.I): self.support_access = "Site-to-Site VPN"
        elif re.search(r"securelink", text, flags=re.I): self.support_access = "SecureLink"
    def set_ha(self, question: str, answer: str):
        q = (question + " " + answer).lower()
        if re.search(r"sql\s+cluster|cluster(ing)?\b|failover cluster|\bag\b", q): self.ha_db = self.ha_db or "cluster supported"
        if re.search(r"application server|app server|web server|application\b", q) and re.search(r"vmware\s*ha", q): self.ha_app = "vmware_ha_only"
    def set_encryption(self, question: str, answer: str):
        text = (question + " " + answer).lower()
        if re.search(r"tls|https|ssl", text): self.encryption_in_transit = self.encryption_in_transit or "TLS/HTTPS"
        if re.search(r"tde|transparent data encryption|sql\s+server\s+tde", text) or re.search(r"aes\s*-?\s*256", text):
            parts = []
            if re.search(r"sql", text): parts.append("SQL Server TDE")
            if re.search(r"aes\s*-?\s*256", text): parts.append("AES-256")
            if re.search(r"key(s)?\s+(managed|provided)\s+by\s+(the\s+)?customer|customer[-\s]?managed", text): parts.append("keys customer-managed")
            detail = "; ".join(parts) if parts else "encrypted at rest"
            if not self.encryption_at_rest or len(detail) > len(self.encryption_at_rest): self.encryption_at_rest = detail

class AutoQLearner:
    def __init__(self):
        self.report = LearnReport()
        self.items: List[LearnItem] = []
        self.tech = TechProfile()

    def _should_reject_answer(self, q: str, a: str) -> Optional[str]:
        qn, an = normtxt(q), normtxt(a)
        if not an: return "empty_answer"
        if looks_like_placeholder(an): return "placeholder_answer"
        if an.lower() == qn.lower(): return "answer_equals_question"
        if is_customer_standard_block(qn): return "customer_standard"
        if re.fullmatch(r"(?i)(yes\s*no|no\s*yes)", an.replace("/", " ")): return "ambiguous_checkbox"
        if len(re.findall(r"yes|no|☐|\[\s*\]", an, flags=re.I)) >= 2 and len(an) < 20: return "checkbox_only"
        if not re.search(r"[A-Za-z]", an): return "non_alpha_answer"
        # At least 3 words OR comma list, unless it's a ports answer
        if len(an.split()) < 3 and ',' not in an and not re.search(r'\bport(s)?\b|\btcp\b|\budp\b', (qn + ' ' + an), flags=re.I):
            return "low_information_answer"
        return None

    def _normalize_and_enrich(self, item: LearnItem):
        q, a = item.question, item.answer
        lower_both = (q + " " + a).lower()
        # Ports only in explicit context
        if re.search(r"\bport(s)?\b|\btcp\b|\budp\b", lower_both) and re.search(r"\d", a):
            self.tech.merge_ports(a)
            item.normalized.setdefault("ports_parsed", True)
        # Latency / RTO
        if re.search(r"latency|ms\b", lower_both): self.tech.set_latency(a)
        if re.search(r"rto|recovery time objective|recovery point objective|rpo", lower_both): self.tech.set_rto(a)
        # Auth / SSO
        if re.search(r"active directory|ldap|saml|imprivata|single sign", lower_both):
            self.tech.set_auth(a or q)
            item.normalized.setdefault("control_stance", {"owner":"customer-managed","compatibility":"product compatible"})
        # Support access
        if re.search(r"vpn|securelink|remote access", lower_both): self.tech.set_support_access(a or q)
        # HA
        if re.search(r"high availability|ha|cluster|\bag\b|failover", lower_both): self.tech.set_ha(q, a)
        # Encryption
        if re.search(r"encrypt|tde|aes|tls|https|ssl", lower_both): self.tech.set_encryption(q, a)

    # DOCX
    def _iter_docx_tables(self, doc) -> Iterable[Tuple[str, str]]:
        for t in doc.tables:
            for r in t.rows:
                cells = [normtxt(c.text) for c in r.cells]
                if len(cells) >= 2:
                    q, a = cells[0], cells[1]
                    if q and a: yield (q, a)
    def _iter_docx_paragraph_pairs(self, doc) -> Iterable[Tuple[str, str]]:
        paras = [normtxt(p.text) for p in doc.paragraphs]
        for i, p in enumerate(paras[:-1]):
            if p.endswith("?") and paras[i+1]:
                yield (p, paras[i+1])
    def process_docx(self, path: str):
        if Document is None:
            self.report.skip("missing_python_docx", path); return
        try: doc = Document(path)
        except Exception as e:
            self.report.skip("docx_open_error", path, {"error": str(e)[:160]}); return
        full_text = "\n".join([p.text for p in doc.paragraphs])
        if is_legal_contract_text(full_text):
            self.report.skip("legal_contract", path); return
        candidates = list(self._iter_docx_tables(doc)) + list(self._iter_docx_paragraph_pairs(doc))
        for (q, a) in candidates:
            if is_customer_standard_block(q):
                self.report.skip("customer_standard", path, {"q": q[:80]}); continue
            rej = self._should_reject_answer(q, a)
            if rej: self.report.skip(rej, path, {"q": q[:80], "a": a[:80]}); continue
            item = LearnItem(source=path, question=q, answer=a)
            self._normalize_and_enrich(item); self.items.append(item); self.report.accept()
        # Doc-level: do not sweep ports; only high-signal
        combined = " \n ".join([normtxt(p.text) for p in doc.paragraphs])
        self.tech.set_latency(combined); self.tech.set_rto(combined)
        self.tech.set_auth(combined); self.tech.set_support_access(combined)
        self.tech.set_ha("", combined); self.tech.set_encryption("", combined)

    # PDF
    def process_pdf(self, path: str):
        if PyPDF2 is None:
            self.report.skip("missing_pypdf2", path); return
        try:
            with open(path,'rb') as f:
                reader = PyPDF2.PdfReader(f); pages_text=[]
                for i in range(len(reader.pages)):
                    try: pages_text.append(reader.pages[i].extract_text() or "")
                    except Exception: pages_text.append("")
        except Exception as e:
            self.report.skip("pdf_open_error", path, {"error": str(e)[:160]}); return
        full_text = "\n".join(pages_text)
        if not full_text.strip(): self.report.skip("empty_pdf", path); return
        if is_legal_contract_text(full_text): self.report.skip("legal_contract", path); return
        if is_blank_survey_pdf(full_text): self.report.skip("blank_form_pdf", path); return
        lines = [l.strip() for l in full_text.splitlines() if l.strip()]
        for i, l in enumerate(lines):
            if re.match(r"^\d{1,3}[\).] ", l) or l.endswith("?"):
                q = l; a = ""
                for j in range(i+1, min(i+6, len(lines))):
                    cand = lines[j]
                    if re.match(r"^\d{1,3}[\).] ", cand) or cand.endswith("?"): break
                    if re.search(r"\bYes\b\s*\bNo\b|☐", cand, flags=re.I): continue
                    a = cand; break
                if not a: continue
                if is_customer_standard_block(q):
                    self.report.skip("customer_standard", path, {"q": q[:80]}); continue
                rej = self._should_reject_answer(q, a)
                if rej: self.report.skip(rej, path, {"q": q[:80], "a": a[:80]}); continue
                item = LearnItem(source=path, question=q, answer=a)
                self._normalize_and_enrich(item); self.items.append(item); self.report.accept()
        self.tech.set_latency(full_text); self.tech.set_rto(full_text)
        self.tech.set_auth(full_text); self.tech.set_support_access(full_text)
        self.tech.set_ha("", full_text); self.tech.set_encryption("", full_text)

    # XLSX
    def process_xlsx(self, path: str):
        if pd is None:
            self.report.skip("missing_pandas", path); return
        try: xls = pd.ExcelFile(path, engine="openpyxl")
        except Exception as e:
            self.report.skip("xlsx_open_error", path, {"error": str(e)[:160]}); return
        allow = re.compile(r"^(solution|security|infrastructure|ms\s*sql\s*server|emr|eMR)$", re.I)
        deny  = re.compile(r"(instructions|use only|dh env|sponsor|vendor|ta process|lawson|pacs|unreviewed)", re.I)
        for sheet in xls.sheet_names:
            if deny.search(sheet):
                self.report.skip("sheet_skipped_policy", path, {"sheet": sheet}); continue
            if not allow.search(sheet):
                self.report.skip("sheet_not_whitelisted", path, {"sheet": sheet}); continue
            try: df = xls.parse(sheet, dtype=str)
            except Exception as e:
                self.report.skip("sheet_parse_error", path, {"sheet": sheet, "error": str(e)[:120]}); continue
            df = df.fillna("")
            cols_low = [c.lower() for c in df.columns]
            q_col = next((df.columns[i] for i,c in enumerate(cols_low) if "question" in c), None)
            a_col = next((df.columns[i] for i,c in enumerate(cols_low) if "response" in c or "vendor response" in c), None)
            if q_col and a_col:
                for _, row in df.iterrows():
                    q, a = normtxt(row[q_col]), normtxt(row[a_col])
                    if not q: continue
                    rej = self._should_reject_answer(q, a)
                    if rej: self.report.skip(rej, path, {"sheet": sheet, "q": q[:60], "a": a[:60]}); continue
                    item = LearnItem(source=f"{path}::{sheet}", question=q, answer=a)
                    self._normalize_and_enrich(item); self.items.append(item); self.report.accept()
            elif df.shape[1] >= 2:
                for _, row in df.iterrows():
                    q, a = normtxt(row.iloc[0]), normtxt(row.iloc[1])
                    if not q: continue
                    rej = self._should_reject_answer(q, a)
                    if rej: self.report.skip(rej, path, {"sheet": sheet, "q": q[:60], "a": a[:60]}); continue
                    item = LearnItem(source=f"{path}::{sheet}", question=q, answer=a)
                    self._normalize_and_enrich(item); self.items.append(item); self.report.accept()

    def process_path(self, path: str):
        ext = os.path.splitext(path)[1].lower()
        if ext == ".docx": self.process_docx(path)
        elif ext == ".pdf": self.process_pdf(path)
        elif ext in (".xlsx",".xlsm"): self.process_xlsx(path)
        else: self.report.skip("unsupported_ext", path, {"ext": ext})

    def run(self, inputs: List[str], outdir: str):
        os.makedirs(outdir, exist_ok=True)
        for inp in inputs:
            if os.path.isdir(inp):
                for root, _, files in os.walk(inp):
                    for fn in files:
                        if os.path.splitext(fn)[1].lower() in (".pdf",".docx",".xlsx",".xlsm"):
                            self.process_path(os.path.join(root, fn))
            else:
                self.process_path(inp)
        with open(os.path.join(outdir,"kb.learned.json"),"w",encoding="utf-8") as f:
            json.dump([item.__dict__ for item in self.items], f, ensure_ascii=False, indent=2)
        with open(os.path.join(outdir,"learn_report.json"),"w",encoding="utf-8") as f:
            json.dump({"accepted": self.report.accepted,
                       "skipped": dict(self.report.skipped),
                       "examples": {k:v for k,v in self.report.examples.items()}},
                      f, ensure_ascii=False, indent=2)
        with open(os.path.join(outdir,"tech_profile.json"),"w",encoding="utf-8") as f:
            json.dump(self.tech.as_dict(), f, ensure_ascii=False, indent=2)

def parse_args(argv: Optional[List[str]] = None) -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Enterprise Q/A learner (Categories 1–5)")
    g = p.add_mutually_exclusive_group(required=True)
    g.add_argument("--in", dest="indir", help="Input directory")
    g.add_argument("--files", nargs="+", help="Explicit list of files to process")
    p.add_argument("--out", dest="outdir", required=True, help="Output directory")
    return p.parse_args(argv)

def main(argv: Optional[List[str]] = None) -> int:
    ns = parse_args(argv)
    learner = AutoQLearner()
    inputs: List[str] = []
    if ns.indir: inputs.append(ns.indir)
    if ns.files: inputs.extend(ns.files)
    learner.run(inputs, ns.outdir)
    print(f"Accepted items: {learner.report.accepted}")
    print(f"Skipped summary: {dict(learner.report.skipped)}")
    print(f"Tech profile written to: {os.path.join(ns.outdir, 'tech_profile.json')}")
    return 0

if __name__ == "__main__":
    sys.exit(main())